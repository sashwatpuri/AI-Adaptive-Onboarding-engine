import fitz  # PyMuPDF
import pdfplumber
from io import BytesIO
from config import config

def parse_pdf(file_bytes: bytes, filename: str = "") -> str:
    """Parse PDF or DOCX files and extract text"""
    
    # Handle DOCX files
    if filename.lower().endswith('.docx'):
        try:
            from docx import Document
            doc = Document(BytesIO(file_bytes))
            
            blocks = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    blocks.append(paragraph.text)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        blocks.append(" | ".join(row_text))
            
            text = "\n".join(blocks)
            print(f"Parsed DOCX: {filename} - {len(text)} chars")
            return clean_text(text)
        except ImportError:
            raise Exception("python-docx library not installed. Please run: pip install python-docx")
        except Exception as e:
            raise Exception(f"Failed to parse DOCX file: {str(e)}")

    # Handle PDF files
    if filename.lower().endswith('.pdf'):
        # Primary for PDF: PyMuPDF
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text("text") + "\n"
            text = text.strip()
            if len(text) > 100:
                print(f"Parsed PDF with PyMuPDF: {filename} - {len(text)} chars")
                return clean_text(text)
        except Exception as e:
            print(f"PyMuPDF failed: {str(e)}, trying pdfplumber...")
        
        # Fallback for PDF: pdfplumber
        try:
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                text = "\n".join(
                    p.extract_text() or "" for p in pdf.pages
                )
            print(f"Parsed PDF with pdfplumber: {filename} - {len(text)} chars")
            return clean_text(text)
        except Exception as e:
            raise Exception(f"Failed to parse PDF with both methods: {str(e)}")
    
    # If no recognized format, raise error
    raise Exception(f"Unsupported file format: {filename}. Supported: .pdf, .docx")

def clean_text(text: str) -> str:
    """Clean and normalize extracted text"""
    import re
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove standalone page numbers
    text = re.sub(r'^\d+$', '', text, flags=re.MULTILINE)
    # Truncate to max chars
    return text[:config.MAX_RESUME_CHARS].strip()
